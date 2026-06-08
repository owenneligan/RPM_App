from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup: A4, standard margins ──────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Default style ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)

def set_font(run, bold=False, italic=False, size=10, underline=False):
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'

def heading_para(doc, text, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_font(run, bold=bold, size=size)
    return p

def body_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_before=2, space_after=2, size=10, indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic, size=size)
    return p

def add_page_break(doc):
    doc.add_page_break()

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_top_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '000000')
    tcBorders.append(top)
    tcPr.append(tcBorders)

def add_bottom_border(cell, double=False):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'double' if double else 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    tcBorders.append(bottom)
    tcPr.append(tcBorders)

def remove_all_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right','insideH','insideV']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                tcBorders.append(el)
            tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10):
    cell.paragraphs[0].clear()
    cell.paragraphs[0].alignment = align
    run = cell.paragraphs[0].add_run(text)
    set_font(run, bold=bold, italic=italic, size=size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trPr.append(trHeight)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run('REGISTERED NUMBER: 13307297 (England and Wales)')
set_font(run, bold=True, size=10)

for _ in range(14):
    doc.add_paragraph()

p = body_para(doc, 'Unaudited Financial Statements for the Year Ended 31 March 2025',
              bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

body_para(doc, 'for', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

body_para(doc, 'NELIBURN LTD', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(0)
run = p.add_run('NELIBURN LTD (REGISTERED NUMBER: 13307297)')
set_font(run, bold=True, size=10)

heading_para(doc, 'Contents of the Financial Statements\nFOR THE YEAR ENDED 31 MARCH 2025', size=10, space_before=8)

t = doc.add_table(rows=5, cols=2)
remove_all_borders(t)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
contents = [
    ('', 'Page'),
    ('Company Information', '1'),
    ('Balance Sheet', '2'),
    ('Notes to the Financial Statements', '4'),
    ("Directors' Report", '6'),
]
for i, (left, right) in enumerate(contents):
    bold = (i == 0)
    cell_text(t.rows[i].cells[0], left, bold=bold)
    cell_text(t.rows[i].cells[1], right, bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER)
set_col_width(t, 0, 12)
set_col_width(t, 1, 4)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — COMPANY INFORMATION  (page 1)
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('NELIBURN LTD')
set_font(run, bold=True, size=10)

heading_para(doc, 'Company Information\nFOR THE YEAR ENDED 31 MARCH 2025', size=10, space_before=4)

doc.add_paragraph()
doc.add_paragraph()

info = [
    ('DIRECTORS:', 'Mr Owen Neligan\nMrs Frances Neligan\nMrs Verity Winterburn'),
    ('REGISTERED OFFICE:', 'Office LG06, 1 Quality Court\nChancery Lane\nLondon\nUnited Kingdom\nWC2A 1HR'),
    ('REGISTERED NUMBER:', '13307297 (England and Wales)'),
    ('ACCOUNTANTS:', 'Silver Arc\nChartered Certified Accountants\n1 Quality Court\nChancery Lane\nLondon\nWC2A 1HR'),
]

for label, value in info:
    t = doc.add_table(rows=1, cols=2)
    remove_all_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_text(t.rows[0].cells[0], label, bold=True)
    # multi-line value
    cell = t.rows[0].cells[1]
    cell.paragraphs[0].clear()
    lines = value.split('\n')
    for idx, line in enumerate(lines):
        if idx == 0:
            run = cell.paragraphs[0].add_run(line)
        else:
            p2 = cell.add_paragraph(line)
        run = cell.paragraphs[idx].add_run('') if idx > 0 else run
        # set font on all paragraphs
    for para in cell.paragraphs:
        for r in para.runs:
            set_font(r, size=10)
        if not para.runs:
            run2 = para.add_run(para.text)
            para.clear()
            run2 = para.add_run(lines[cell.paragraphs.index(para)] if cell.paragraphs.index(para) < len(lines) else '')
            set_font(run2, size=10)
    set_col_width(t, 0, 5)
    set_col_width(t, 1, 11)
    doc.add_paragraph()

# Page number note
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('Page 1')
set_font(run, size=9)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — BALANCE SHEET  (page 2)
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('NELIBURN LTD (REGISTERED NUMBER: 13307297)')
set_font(run, bold=True, size=10)

heading_para(doc, 'Balance Sheet\n31 MARCH 2025', size=10, space_before=4)

# Build balance sheet table: cols = [description, notes, 2025 sub, 2025 total, 2024 sub, 2024 total]
# We'll use 6 columns
bs_rows = [
    # (description, notes, sub25, tot25, sub24, tot24, bold, top_border_tot, bottom_border_tot, double_bottom)
    ('', 'Notes', '', '31/3/25\n£', '', '31/3/24\n£', True, False, False, False),
    ('FIXED ASSETS', '', '', '', '', '', True, False, False, False),
    ('Investment property', '4', '', '225,000', '', '225,000', False, False, False, False),
    ('', '', '', '', '', '', False, False, False, False),
    ('CURRENT ASSETS', '', '', '', '', '', True, False, False, False),
    ('Debtors', '5', '3,874', '', '4,915', '', False, False, False, False),
    ('Cash at bank and in hand', '', '3,796', '', '27,110', '', False, False, False, False),
    ('', '', '7,670', '', '32,025', '', False, True, False, False),
    ('', '', '', '', '', '', False, False, False, False),
    ('CREDITORS: Amounts falling due within one year', '6', '(68,436)', '', '(79,114)', '', False, False, False, False),
    ('NET CURRENT LIABILITIES', '', '', '(60,766)', '', '(47,089)', True, False, False, False),
    ('TOTAL ASSETS LESS CURRENT LIABILITIES', '', '', '164,234', '', '177,911', True, True, True, False),
    ('', '', '', '', '', '', False, False, False, False),
    ('CREDITORS: Amounts falling due after more than one year', '7', '', '(173,791)', '', '(173,791)', False, False, False, False),
    ('', '', '', '', '', '', False, False, False, False),
    ('NET LIABILITIES', '', '', '(9,557)', '', '4,120', True, True, True, False),
    ('', '', '', '', '', '', False, False, False, False),
    ('CAPITAL AND RESERVES', '', '', '', '', '', True, False, False, False),
    ('Called up share capital', '', '', '3', '', '3', False, False, False, False),
    ('Fair value reserve', '8', '', '22,546', '', '22,546', False, False, False, False),
    ('Retained earnings', '', '', '(32,106)', '', '(18,429)', False, False, False, False),
    ('TOTAL EQUITY', '', '', '(9,557)', '', '4,120', True, True, True, True),
]

t = doc.add_table(rows=len(bs_rows), cols=6)
remove_all_borders(t)
t.alignment = WD_TABLE_ALIGNMENT.LEFT

col_widths = [7.5, 1.0, 2.0, 2.0, 2.0, 2.0]
for ci, w in enumerate(col_widths):
    set_col_width(t, ci, w)

for ri, row_data in enumerate(bs_rows):
    desc, notes, sub25, tot25, sub24, tot24, bold, top_tot, bot_tot, dbl = row_data
    cells = t.rows[ri].cells
    cell_text(cells[0], desc, bold=bold, size=10)
    cell_text(cells[1], notes, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    cell_text(cells[2], sub25, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    cell_text(cells[3], tot25, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    cell_text(cells[4], sub24, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    cell_text(cells[5], tot24, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    if top_tot:
        add_top_border(cells[3])
        add_top_border(cells[5])
    if bot_tot:
        add_bottom_border(cells[3], double=dbl)
        add_bottom_border(cells[5], double=dbl)

doc.add_paragraph()

# Statutory statements
statutory = [
    'The company is entitled to exemption from audit under Section 477 of the Companies Act 2006 for the year ended 31 March 2025.',
    '',
    'The members have not required the company to obtain an audit of its financial statements for the year ended 31 March 2025 in accordance with Section 476 of the Companies Act 2006.',
    '',
    'The directors acknowledge their responsibilities for:',
    '(a)  ensuring that the company keeps accounting records which comply with Sections 386 and 387 of the Companies Act 2006; and',
    '(b)  preparing financial statements which give a true and fair view of the state of affairs of the company as at the end of each financial year and of its profit or loss for each financial year in accordance with the requirements of Sections 394 and 395 and which otherwise comply with the requirements of the Companies Act 2006 relating to financial statements, so far as applicable to the company.',
]
for s in statutory:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if s.startswith('('):
        p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(s)
    set_font(run, size=9)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run('The financial statements have been prepared and delivered in accordance with the provisions applicable to companies subject to the small companies regime.')
set_font(run, size=9)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run = p.add_run('In accordance with Section 444 of the Companies Act 2006, the Income Statement has not been delivered.')
set_font(run, size=9)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run = p.add_run('The financial statements were approved by the Board of Directors and authorised for issue on _________________________ and were signed on its behalf by:')
set_font(run, size=9)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Mr Owen Neligan  –  Director')
set_font(run, size=10)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run('Mrs Frances Neligan  –  Director')
set_font(run, size=10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run('The notes form part of these financial statements')
set_font(run, italic=True, size=9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Page 2')
set_font(run, size=9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('continued...')
set_font(run, italic=True, size=9)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 5 — BALANCE SHEET continued  (page 3)
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('NELIBURN LTD (REGISTERED NUMBER: 13307297)')
set_font(run, bold=True, size=10)

heading_para(doc, 'Balance Sheet – continued\n31 MARCH 2025', size=10, space_before=4)

cont_paras = [
    'These financial statements have been prepared and delivered in accordance with the provisions applicable to companies subject to the small companies regime.',
    'In accordance with Section 444 of the Companies Act 2006, the Income Statement has not been delivered.',
]
for cp in cont_paras:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(cp)
    set_font(run, size=10)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run('The financial statements were approved by the Board of Directors and authorised for issue on _________________________ and were signed on its behalf by:')
set_font(run, size=10)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Mr Owen Neligan  –  Director')
set_font(run, size=10)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
run = p.add_run('Mrs Frances Neligan  –  Director')
set_font(run, size=10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('The notes form part of these financial statements')
set_font(run, italic=True, size=9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Page 3')
set_font(run, size=9)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 6 — NOTES TO THE FINANCIAL STATEMENTS  (page 4)
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('NELIBURN LTD (REGISTERED NUMBER: 13307297)')
set_font(run, bold=True, size=10)

heading_para(doc, 'Notes to the Financial Statements\nFOR THE YEAR ENDED 31 MARCH 2025', size=10, space_before=4)

# Note 1
heading_para(doc, '1.     STATUTORY INFORMATION', size=10, space_before=8, space_after=2)
body_para(doc, 'Neliburn Ltd is a private company limited by shares, registered in England and Wales (number 13307297). The registered office is Office LG06, 1 Quality Court, Chancery Lane, London, WC2A 1HR.', size=10, space_before=2, space_after=4)

# Note 2
heading_para(doc, '2.     ACCOUNTING POLICIES', size=10, space_before=8, space_after=2)

policies = [
    ('Basis of preparing the financial statements',
     'These financial statements have been prepared in accordance with Financial Reporting Standard 102 "The Financial Reporting Standard applicable in the UK and Republic of Ireland", including the provisions of Section 1A "Small Entities", and the Companies Act 2006. The financial statements have been prepared under the historical cost convention as modified by the revaluation of the investment property.'),
    ('Turnover',
     'Turnover represents rental income receivable in respect of the period, measured at the gross contracted rent per the tenancy agreement. The company is not registered for VAT. Agent management fees are presented as an expense rather than netted against income.'),
    ('Investment property',
     'Investment property is initially recognised at cost and is subsequently carried at fair value at each balance sheet date, with changes in fair value recognised in the profit and loss account. At 31 March 2025 no formal independent revaluation was commissioned; the directors are satisfied that the carrying value of £225,000, being the valuation determined by external valuers on 29 December 2023, remains a reasonable approximation of fair value at the balance sheet date.'),
    ('Taxation',
     'Corporation Tax is recognised on taxable profits at rates enacted at the balance sheet date. As the company is loss-making in the period, no current tax charge arises. Deferred tax is recognised only where it is probable that a timing difference will reverse.'),
    ('Cash and cash equivalents',
     'Cash and cash equivalents comprise cash held in bank accounts. There are no short-term investments.'),
    ('Financial instruments',
     'Basic financial instruments, including trade debtors, trade creditors, and director loan accounts, are recognised at transaction price. The bank loan is carried at the amount outstanding. Interest payable on the bank loan is recognised on an accruals basis.'),
]

for title, text in policies:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.7)
    run = p.add_run(title)
    set_font(run, bold=True, size=10)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after  = Pt(3)
    p2.paragraph_format.left_indent  = Cm(0.7)
    run2 = p2.add_run(text)
    set_font(run2, size=10)

# Note 3
heading_para(doc, '3.     EMPLOYEES AND DIRECTORS', size=10, space_before=8, space_after=2)
body_para(doc, 'The average number of employees during the year was NIL (2024 – NIL). No directors’ remuneration was paid during the year (2024 – NIL).', size=10, space_before=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
run = p.add_run('The notes form part of these financial statements')
set_font(run, italic=True, size=9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Page 4')
set_font(run, size=9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('continued...')
set_font(run, italic=True, size=9)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 7 — NOTES continued  (page 5)
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('NELIBURN LTD (REGISTERED NUMBER: 13307297)')
set_font(run, bold=True, size=10)

heading_para(doc, 'Notes to the Financial Statements – continued\nFOR THE YEAR ENDED 31 MARCH 2025', size=10, space_before=4)

# Note 4 — Investment Property
heading_para(doc, '4.     INVESTMENT PROPERTY', size=10, space_before=8, space_after=4)

ip_rows = [
    ('', 'Total £', True),
    ('FAIR VALUE', '', True),
    ('At 1 April 2024', '225,000', False),
    ('Revaluations in the year', '—', False),
    ('At 31 March 2025', '225,000', True),
    ('NET BOOK VALUE', '', True),
    ('At 31 March 2025', '225,000', True),
    ('At 31 March 2024', '225,000', True),
]

t = doc.add_table(rows=len(ip_rows), cols=2)
remove_all_borders(t)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
set_col_width(t, 0, 11)
set_col_width(t, 1, 5)

for ri, (desc, val, bold) in enumerate(ip_rows):
    cell_text(t.rows[ri].cells[0], desc, bold=bold, size=10)
    cell_text(t.rows[ri].cells[1], val, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    if ri == 4:
        add_top_border(t.rows[ri].cells[1])
    if ri in (6, 7):
        add_bottom_border(t.rows[ri].cells[1], double=(ri == 7))

body_para(doc, 'Investment property was independently valued on an open market basis on 29 December 2023 by external valuers. No formal revaluation was carried out at 31 March 2025; the directors consider the carrying value to represent a reasonable approximation of fair value at the balance sheet date.', size=10, space_before=6)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run = p.add_run('If investment property had not been revalued it would have been included at the following historical cost:')
set_font(run, size=10)

hc_rows = [
    ('', '31/3/25\n£', '31/3/24\n£'),
    ('Cost', '202,454', '202,454'),
]
t2 = doc.add_table(rows=2, cols=3)
remove_all_borders(t2)
t2.alignment = WD_TABLE_ALIGNMENT.LEFT
set_col_width(t2, 0, 8)
set_col_width(t2, 1, 4)
set_col_width(t2, 2, 4)
for ri, (d, v1, v2) in enumerate(hc_rows):
    bold = (ri == 0)
    cell_text(t2.rows[ri].cells[0], d, bold=bold, size=10)
    cell_text(t2.rows[ri].cells[1], v1, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    cell_text(t2.rows[ri].cells[2], v2, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    if ri == 1:
        add_bottom_border(t2.rows[ri].cells[1])
        add_bottom_border(t2.rows[ri].cells[2])

# Note 5 — Debtors
heading_para(doc, '5.     DEBTORS: AMOUNTS FALLING DUE WITHIN ONE YEAR', size=10, space_before=10, space_after=4)

deb_rows = [
    ('', '31/3/25\n£', '31/3/24\n£', True),
    ('Other debtors', '3,874', '4,915', False),
]
t3 = doc.add_table(rows=2, cols=3)
remove_all_borders(t3)
t3.alignment = WD_TABLE_ALIGNMENT.LEFT
set_col_width(t3, 0, 8)
set_col_width(t3, 1, 4)
set_col_width(t3, 2, 4)
for ri, (d, v1, v2, bold) in enumerate(deb_rows):
    cell_text(t3.rows[ri].cells[0], d, bold=bold, size=10)
    cell_text(t3.rows[ri].cells[1], v1, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    cell_text(t3.rows[ri].cells[2], v2, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    if ri == 1:
        add_bottom_border(t3.rows[ri].cells[1])
        add_bottom_border(t3.rows[ri].cells[2])

# Note 6 — Creditors <1yr
heading_para(doc, '6.     CREDITORS: AMOUNTS FALLING DUE WITHIN ONE YEAR', size=10, space_before=10, space_after=4)

cred_rows = [
    ('', '31/3/25\n£', '31/3/24\n£', True),
    ('Trade creditors', '—', '1,080', False),
    ('Director loan accounts', '68,436', '78,034', False),
    ('Total', '68,436', '79,114', True),
]
t4 = doc.add_table(rows=4, cols=3)
remove_all_borders(t4)
t4.alignment = WD_TABLE_ALIGNMENT.LEFT
set_col_width(t4, 0, 8)
set_col_width(t4, 1, 4)
set_col_width(t4, 2, 4)
for ri, (d, v1, v2, bold) in enumerate(cred_rows):
    cell_text(t4.rows[ri].cells[0], d, bold=bold, size=10)
    cell_text(t4.rows[ri].cells[1], v1, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    cell_text(t4.rows[ri].cells[2], v2, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    if ri == 2:
        add_top_border(t4.rows[ri].cells[1])
        add_top_border(t4.rows[ri].cells[2])
    if ri == 3:
        add_bottom_border(t4.rows[ri].cells[1], double=True)
        add_bottom_border(t4.rows[ri].cells[2], double=True)

body_para(doc, 'The director loan accounts represent monies lent to the company by the directors. The amounts are unsecured, interest-free and repayable on demand. During the year ended 31 March 2025, the company repaid £4,565 to Mr Owen Neligan and £4,000 to Mrs Frances Neligan.', size=10, space_before=6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
run = p.add_run('The notes form part of these financial statements')
set_font(run, italic=True, size=9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Page 5')
set_font(run, size=9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('continued...')
set_font(run, italic=True, size=9)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 8 — NOTES continued  (page 6)
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('NELIBURN LTD (REGISTERED NUMBER: 13307297)')
set_font(run, bold=True, size=10)

heading_para(doc, 'Notes to the Financial Statements – continued\nFOR THE YEAR ENDED 31 MARCH 2025', size=10, space_before=4)

# Note 7 — Creditors >1yr
heading_para(doc, '7.     CREDITORS: AMOUNTS FALLING DUE AFTER MORE THAN ONE YEAR', size=10, space_before=8, space_after=4)

lt_rows = [
    ('', '31/3/25\n£', '31/3/24\n£', True),
    ('Bank loans', '173,791', '173,791', False),
    ('', '', '', False),
    ('Amounts falling due in more than five years:', '', '', False),
    ('Repayable by instalments', '', '', False),
    ('Bank loans more than 5 years by instalment', '173,791', '173,791', False),
]
t5 = doc.add_table(rows=len(lt_rows), cols=3)
remove_all_borders(t5)
t5.alignment = WD_TABLE_ALIGNMENT.LEFT
set_col_width(t5, 0, 8)
set_col_width(t5, 1, 4)
set_col_width(t5, 2, 4)
for ri, (d, v1, v2, bold) in enumerate(lt_rows):
    cell_text(t5.rows[ri].cells[0], d, bold=bold, size=10)
    cell_text(t5.rows[ri].cells[1], v1, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    cell_text(t5.rows[ri].cells[2], v2, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    if ri == 5:
        add_top_border(t5.rows[ri].cells[1])
        add_top_border(t5.rows[ri].cells[2])
        add_bottom_border(t5.rows[ri].cells[1])
        add_bottom_border(t5.rows[ri].cells[2])

body_para(doc, 'The bank loan relates to a buy-to-let mortgage with Fleet Mortgages Ltd (account number 1100082162), secured by a first legal charge over the company\'s investment property at 11 Hatfield Walk, York, YO24 3LX. The mortgage is interest only at a variable rate of 5.44% per annum (effective from 1 April 2024). The remaining term at the balance sheet date is approximately 26 years and 10 months. Total interest charged in the year amounted to £9,455 (2024: not separately disclosed).', size=10, space_before=6)

# Note 8 — Reserves
heading_para(doc, '8.     RESERVES', size=10, space_before=10, space_after=4)

res_rows = [
    ('', 'Fair value\nreserve £', True),
    ('Non-distributable reserve', '', True),
    ('At 1 April 2024', '22,546', False),
    ('Movement in year', '—', False),
    ('At 31 March 2025', '22,546', True),
]
t6 = doc.add_table(rows=5, cols=2)
remove_all_borders(t6)
t6.alignment = WD_TABLE_ALIGNMENT.LEFT
set_col_width(t6, 0, 10)
set_col_width(t6, 1, 6)
for ri, (d, v, bold) in enumerate(res_rows):
    cell_text(t6.rows[ri].cells[0], d, bold=bold, size=10)
    cell_text(t6.rows[ri].cells[1], v, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    if ri == 4:
        add_top_border(t6.rows[ri].cells[1])
        add_bottom_border(t6.rows[ri].cells[1])

body_para(doc, 'The fair value reserve represents the cumulative uplift of the investment property from historical cost (£202,454) to fair value (£225,000), arising on revaluation in the year ended 31 March 2024. This reserve is non-distributable.', size=10, space_before=6)

# Note 9 — Related party
heading_para(doc, '9.     RELATED PARTY TRANSACTIONS', size=10, space_before=10, space_after=4)
body_para(doc, 'The directors are the beneficial owners of the company. Director loan account balances are disclosed in Note 6. The amounts are unsecured, interest-free and repayable on demand. No other related party transactions requiring disclosure under FRS 102 Section 1A have been identified during the year.', size=10, space_before=2)

# Note 10
heading_para(doc, '10.    POST BALANCE SHEET EVENTS', size=10, space_before=10, space_after=4)
body_para(doc, 'There are no post balance sheet events requiring disclosure.', size=10, space_before=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
run = p.add_run('The notes form part of these financial statements')
set_font(run, italic=True, size=9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Page 6')
set_font(run, size=9)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 9 — DIRECTORS' REPORT
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('NELIBURN LTD (REGISTERED NUMBER: 13307297)')
set_font(run, bold=True, size=10)

heading_para(doc, "Directors' Report\nFOR THE YEAR ENDED 31 MARCH 2025", size=10, space_before=4)

dr_text = (
    "The directors present their report for the year ended 31 March 2025. "
    "The company's principal activity during the year was the holding and letting of a single residential investment property at "
    "11 Hatfield Walk, York, YO24 3LX, which was subject to a programme of renovation works during the period April to June 2024 "
    "before being let to tenants through Ashtons Lettings & Management from July 2024. "
    "The company incurred a loss before taxation of £13,677 for the year (2024: loss of £11,488), principally reflecting renovation "
    "expenditure during the void period and the full year's mortgage interest charge of £9,455. "
    "The directors do not recommend the payment of a dividend. "
    "The directors who served during the year were Mr Owen Neligan, Mrs Frances Neligan and Mrs Verity Winterburn."
)
body_para(doc, dr_text, size=10, space_before=8)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Signed on behalf of the Board:')
set_font(run, size=10)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Mr Owen Neligan  –  Director')
set_font(run, size=10)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
run = p.add_run('Date: _______________________')
set_font(run, size=10)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run = p.add_run('Registered office: Office LG06, 1 Quality Court, Chancery Lane, London, WC2A 1HR')
set_font(run, size=10)

# ── Save ──────────────────────────────────────────────────────────────────────
path = '/home/user/RPM_App/Neliburn_Ltd_Accounts_YE_31March2025.docx'
doc.save(path)
print(f'Saved: {path}')
